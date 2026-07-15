"""
Build a downloadable ZIP of a completed run: every agent's output as readable
Markdown, the key generated code extracted into real files, and a beginner-
friendly HOW-TO-HOST.docx so a non-technical user can deploy what they built.
"""
from __future__ import annotations

import io
import re
import zipfile
from typing import Any

from docx import Document
from docx.shared import Pt


# agent id -> human-readable Markdown filename in the zip root.
DOC_FILES: dict[str, str] = {
    "pm": "01-PRODUCT-REQUIREMENTS.md",
    "database": "02-DATABASE.md",
    "backend": "03-BACKEND.md",
    "uiux": "04-UI-UX-DESIGN.md",
    "frontend": "05-FRONTEND.md",
    "reviewer": "06-CODE-REVIEW.md",
    "qa": "07-QA-TEST-PLAN.md",
    "security": "08-SECURITY-REVIEW.md",
    "remediation": "09-FIXES-APPLIED.md",
    "devops": "10-DEPLOYMENT-PLAN.md",
    "docs": "README.md",
}


def _code_blocks(text: str, langs: tuple[str, ...] | None = None) -> list[str]:
    """Return the bodies of fenced code blocks, optionally filtered by language."""
    if not text:
        return []
    out: list[str] = []
    for lang, body in re.findall(r"```([\w+.-]*)\s*\n(.*?)```", text, re.S):
        if langs is None or lang.lower() in langs:
            out.append(body.strip("\n"))
    return out


def _first(blocks: list[str]) -> str | None:
    return blocks[0] if blocks else None


def _extract_code_files(outputs: dict[str, str]) -> dict[str, str]:
    """Pull the most useful generated code into real, named files."""
    files: dict[str, str] = {}

    frontend = outputs.get("frontend", "") or ""
    remediation = outputs.get("remediation", "") or ""
    database = outputs.get("database", "") or ""
    backend = outputs.get("backend", "") or ""
    devops = outputs.get("devops", "") or ""

    # Frontend: prefer the remediated (fixed) component; else the original.
    app = _first(_code_blocks(remediation, ("tsx", "jsx"))) or _first(
        _code_blocks(frontend, ("tsx", "jsx"))
    )
    if app:
        files["frontend/App.tsx"] = app
    mockup = _first(_code_blocks(frontend, ("html",)))
    if mockup:
        files["frontend/mockup.html"] = mockup

    schema = _first(_code_blocks(database, ("sql",)))
    if schema:
        files["database/schema.sql"] = schema

    api = _first(_code_blocks(backend, ("python", "py")))
    if api:
        files["backend/app.py"] = (
            "# Starting point extracted from the generated backend design.\n"
            "# Review before running in production.\n\n" + api
        )

    dockerfile = _first(_code_blocks(devops, ("dockerfile", "docker")))
    if dockerfile:
        files["backend/Dockerfile"] = dockerfile

    return files


def _hosting_docx(idea: str, outputs: dict[str, str]) -> bytes:
    """A friendly, step-by-step hosting guide as a .docx."""
    doc = Document()
    doc.add_heading("How to Host Your Product", level=0)

    intro = doc.add_paragraph()
    intro.add_run("Your idea: ").bold = True
    intro.add_run(idea or "your product")

    doc.add_paragraph(
        "This guide walks you through putting your app online so anyone can use it. "
        "You do not need to be a developer. Everything below uses free plans. Take it "
        "one step at a time — each step says exactly what to click."
    )

    doc.add_heading("What you have in this folder", level=1)
    for line in (
        "frontend/ — the website people see (App.tsx and a mockup.html preview).",
        "backend/ — the server that powers it (app.py and a Dockerfile).",
        "database/schema.sql — the database tables to create.",
        "The numbered .md files — the full plan: requirements, design, tests, "
        "security review, and the fixes that were applied.",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("What you'll need (all free)", level=1)
    for line in (
        "A GitHub account — to store your code.",
        "A Supabase account — the database.",
        "A Render account — runs the backend server.",
        "A Vercel account — hosts the website.",
    ):
        doc.add_paragraph(line, style="List Bullet")

    steps = [
        (
            "Put the code on GitHub",
            "Create a new repository on github.com and upload this whole folder to it. "
            "This is where Vercel and Render will read your code from.",
        ),
        (
            "Create the database (Supabase)",
            "Go to supabase.com and create a project. Open the SQL Editor, paste the "
            "contents of database/schema.sql, and run it to create your tables. "
            "In Project Settings > API, copy the Project URL and the two keys "
            "(anon and service_role) — you'll paste them in the next steps.",
        ),
        (
            "Host the backend (Render)",
            "On render.com, choose New > Web Service and pick your GitHub repo. Set the "
            "root directory to backend/. Add your keys as environment variables (the "
            "backend/.env example lists which ones). When it finishes, copy the "
            "backend's URL — it ends in .onrender.com.",
        ),
        (
            "Host the website (Vercel)",
            "On vercel.com, choose Add New > Project and pick the same GitHub repo. Set "
            "the backend URL and your Supabase URL + anon key as environment variables. "
            "Deploy. Vercel gives you a public link — that's your live site.",
        ),
        (
            "Connect the two",
            "Back in Render, set the FRONTEND_ORIGIN variable to your Vercel link so the "
            "website is allowed to talk to the backend. Save and let it redeploy.",
        ),
        (
            "Test it",
            "Open your Vercel link and try the main actions. If something can't reach the "
            "backend, double-check the backend URL in Vercel and FRONTEND_ORIGIN in "
            "Render match your live addresses exactly.",
        ),
    ]
    doc.add_heading("Step-by-step", level=1)
    for i, (title, body) in enumerate(steps, 1):
        h = doc.add_paragraph()
        run = h.add_run(f"Step {i}: {title}")
        run.bold = True
        run.font.size = Pt(13)
        doc.add_paragraph(body)

    doc.add_heading("If you get stuck", level=1)
    for line in (
        "Backend link shows 'Not Found' — that's normal; the backend has no homepage. "
        "Add /health to the URL to check it's alive.",
        "First load is slow — free backends sleep when idle and take ~50 seconds to "
        "wake up on the first visit.",
        "Buttons can't reach the server — the FRONTEND_ORIGIN in Render must match your "
        "Vercel link exactly (no trailing slash).",
    ):
        doc.add_paragraph(line, style="List Bullet")

    devops = outputs.get("devops", "").strip()
    if devops:
        doc.add_heading("Appendix: detailed deployment plan (generated)", level=1)
        for para in devops.splitlines():
            doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_zip(idea: str, outputs: dict[str, Any]) -> bytes:
    """Assemble the full downloadable bundle and return the zip as bytes."""
    outputs = {k: (v if isinstance(v, str) else str(v)) for k, v in (outputs or {}).items()}

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        # Beginner hosting guide.
        z.writestr("HOW-TO-HOST.docx", _hosting_docx(idea, outputs))

        # Every agent's output as readable Markdown.
        for agent_id, filename in DOC_FILES.items():
            text = outputs.get(agent_id)
            if text:
                z.writestr(filename, text)

        # Key generated code as real files.
        for path, content in _extract_code_files(outputs).items():
            z.writestr(path, content)

        # A short index so the folder isn't a mystery.
        z.writestr(
            "START-HERE.txt",
            "Open HOW-TO-HOST.docx first — it explains how to put your product "
            "online, step by step.\n\nEverything else: the numbered .md files are the "
            "full plan; the frontend/, backend/ and database/ folders hold the "
            "generated code.\n",
        )

    return buf.getvalue()
